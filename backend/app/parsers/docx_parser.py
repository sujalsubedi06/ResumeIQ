import io
import docx
from app.core.exceptions import ParseFailedError
from app.core.logging import logger
from app.parsers.base import BaseParser
from app.schemas.parser import ParsedDocument, ResumeMetadata

class DOCXParser(BaseParser):
    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        try:
            stream = io.BytesIO(file_bytes)
            doc = docx.Document(stream)

            text_parts = []
            for paragraph in doc.paragraphs:
                p_text = paragraph.text.strip()
                if p_text:
                    text_parts.append(p_text)

            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            full_text = "\n".join(text_parts)
            word_count = len(full_text.split())

            metadata = ResumeMetadata(
                fileName=filename,
                fileType="docx",
                pageCount=None,  # DOCX format does not have explicit pagination in raw XML
                wordCount=word_count,
                fileSizeBytes=len(file_bytes),
            )

            return ParsedDocument(text=full_text, metadata=metadata)
        except Exception as e:
            logger.error(f"DOCX parsing error for file {filename}: {str(e)}")
            raise ParseFailedError(f"Failed to parse DOCX document '{filename}': {str(e)}")
