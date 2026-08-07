import io
import docx
import pytest
from app.core.exceptions import ParseFailedError
from app.parsers.docx_parser import DOCXParser

def create_sample_docx_bytes(content: str = "Jane Developer\nFull Stack Engineer") -> bytes:
    doc = docx.Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    
    # Add a sample table
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Skill"
    hdr_cells[1].text = "Level"

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()

def test_docx_parser_success():
    parser = DOCXParser()
    sample_bytes = create_sample_docx_bytes("Jane Developer\nFull Stack Engineer\nReact, TypeScript")
    result = parser.parse(sample_bytes, "resume.docx")

    assert result.metadata.fileName == "resume.docx"
    assert result.metadata.fileType == "docx"
    assert result.metadata.pageCount is None
    assert "Jane Developer" in result.text
    assert "Full Stack Engineer" in result.text
    assert "Skill | Level" in result.text
    assert result.metadata.wordCount > 0

def test_docx_parser_invalid_bytes():
    parser = DOCXParser()
    invalid_bytes = b"Corrupted docx content"
    with pytest.raises(ParseFailedError) as exc_info:
        parser.parse(invalid_bytes, "corrupt.docx")
    assert "Failed to parse DOCX document" in str(exc_info.value)
